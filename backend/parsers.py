"""文档格式解析：扩展名校验 + 文本提取（纯函数，无外部依赖）。"""

from pathlib import Path
from io import BytesIO
from zipfile import BadZipFile
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from docx import Document

ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def is_allowed_extension(filename: str) -> bool:
    """文件名后缀是否属于支持格式（忽略大小写，兼容 .PDF）。"""
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def extract_text(filename: str, content: bytes) -> str:
    """按扩展名把原始字节提取为纯文本；上传与文档预览共用。"""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        doc = Document(BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    return content.decode("utf-8")


def render_pdf_pages(content: bytes) -> list:
    """把 PDF 每页渲染成 PNG 字节，供扫描件走视觉模型识别。"""
    import pymupdf

    doc = pymupdf.open(stream=content, filetype="pdf")
    pages = []
    for page in doc:
        pix = page.get_pixmap(dpi=200)
        pages.append(pix.tobytes("png"))
    return pages


def extract_embedded_images(filename: str, content: bytes) -> list:
    """提取文档内嵌图片字节（PDF 用 pymupdf，DOCX 用 zipfile），无则返回 []。"""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        import pymupdf

        doc = pymupdf.open(stream=content, filetype="pdf")
        xrefs = set()
        for page in doc:
            for img in page.get_images(full=True):
                xrefs.add(img[0])  # 同一图被多页引用会重复，按 xref 去重
        return [doc.extract_image(x)["image"] for x in xrefs]
    if suffix == ".docx":
        import zipfile

        with zipfile.ZipFile(BytesIO(content)) as z:
            return [
                z.read(n)
                for n in z.namelist()
                if n.startswith("word/media/") and not n.endswith("/")
            ]
    return []
