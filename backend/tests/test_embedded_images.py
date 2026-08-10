"""文档嵌入图片 → 视觉识别进库（_extract_document_text）测试。"""

import io

import llm
import main
import parsers


def _make_png_bytes(text="IMG-888") -> bytes:
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (300, 80), "white")
    ImageDraw.Draw(img).text((10, 30), text, fill="black")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _make_docx_with_image() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph("这是文档正文。")
    doc.add_picture(io.BytesIO(_make_png_bytes()))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf_with_image() -> bytes:
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "TEXT PDF 456")
    page.insert_image(page.rect, stream=_make_png_bytes())
    return doc.tobytes()


def _make_docx_text_only() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph("纯文字文档。")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_with_image_appends_vision_text(monkeypatch):
    monkeypatch.setattr(llm, "vision_extract_text", lambda pages: "图里是 IMG-888")
    text = main._extract_document_text("doc.docx", _make_docx_with_image())
    assert "这是文档正文" in text
    assert "IMG-888" in text


def test_pdf_with_text_and_image_appends_vision(monkeypatch):
    monkeypatch.setattr(llm, "vision_extract_text", lambda pages: "PDF 图内容 456")
    text = main._extract_document_text("doc.pdf", _make_pdf_with_image())
    assert "TEXT PDF 456" in text
    assert "PDF 图内容 456" in text


def test_text_only_docx_does_not_call_vision(monkeypatch):
    called = []

    def fake_vision(pages):
        called.append(True)
        return "不应被调用"

    monkeypatch.setattr(llm, "vision_extract_text", fake_vision)
    text = main._extract_document_text("doc.docx", _make_docx_text_only())
    assert "纯文字文档" in text
    assert not called  # 没图，不该触发视觉


def test_pure_image_docx_uses_vision(monkeypatch):
    from docx import Document

    doc = Document()  # 无段落，只有一张图
    doc.add_picture(io.BytesIO(_make_png_bytes()))
    buf = io.BytesIO()
    doc.save(buf)

    monkeypatch.setattr(llm, "vision_extract_text", lambda pages: "只有图的文档内容")
    text = main._extract_document_text("image_only.docx", buf.getvalue())
    assert "只有图的文档内容" in text


def test_extract_embedded_images_counts():
    assert len(parsers.extract_embedded_images("doc.docx", _make_docx_with_image())) == 1
    assert len(parsers.extract_embedded_images("doc.pdf", _make_pdf_with_image())) == 1
    assert parsers.extract_embedded_images("doc.docx", _make_docx_text_only()) == []
    assert parsers.extract_embedded_images("doc.txt", b"hello") == []
