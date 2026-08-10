"""extract_text / is_allowed_extension 纯函数测试，不依赖 DB 或网络。"""

import io
from zipfile import BadZipFile

import pytest
from pypdf.errors import PdfReadError

import parsers


def _make_pdf(text: str) -> bytes:
    """手搓一个最小合法 PDF，含一行文本。"""
    content = f"BT /F1 24 Tf 72 720 Td ({text}) Tj ET".encode("latin-1")
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length %d >>\nstream\n" % len(content) + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objs, 1):
        offsets.append(out.tell())
        out.write(b"%d 0 obj\n" % i + obj + b"\nendobj\n")
    xref = out.tell()
    out.write(b"xref\n0 %d\n" % 6 + b"0000000000 65535 f \n")
    for off in offsets:
        out.write(b"%010d 00000 n \n" % off)
    out.write(b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n" % xref)
    return out.getvalue()


def _make_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph("SmartRAG 支持 Word 文档。")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "黄栩晖"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_txt_extract():
    assert parsers.extract_text("a.txt", "你好 世界".encode("utf-8")) == "你好 世界"


def test_md_extract():
    assert parsers.extract_text("guide.md", "# 标题\n正文".encode("utf-8")) == "# 标题\n正文"


def test_uppercase_extension_pdf():
    assert parsers.extract_text("A.PDF", _make_pdf("hello")) == "hello"


def test_pdf_extract():
    assert "ZETA-7" in parsers.extract_text("sample.pdf", _make_pdf("SmartRAG code is ZETA-7"))


def test_docx_extract_paragraph_and_table():
    text = parsers.extract_text("sample.docx", _make_docx())
    assert "SmartRAG 支持 Word 文档。" in text
    assert "姓名 | 黄栩晖" in text


def test_corrupt_pdf_raises():
    with pytest.raises(PdfReadError):
        parsers.extract_text("bad.pdf", b"not a pdf")


def test_bad_docx_raises():
    with pytest.raises(BadZipFile):
        parsers.extract_text("bad.docx", b"not a zip")


def test_is_allowed_extension():
    assert parsers.is_allowed_extension("a.txt")
    assert parsers.is_allowed_extension("guide.md")
    assert parsers.is_allowed_extension("README.PDF")
    assert parsers.is_allowed_extension("resume.docx")
    assert not parsers.is_allowed_extension("evil.exe")
    assert not parsers.is_allowed_extension("conversations.db")
    assert not parsers.is_allowed_extension(".env")
